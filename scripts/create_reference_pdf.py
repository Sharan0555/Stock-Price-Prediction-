from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, StyleSheet1, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    LongTable,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.platypus.tables import TableStyle


SOURCE_DOCX = Path("/Users/sharanpatil/Downloads/Stock_Price_Prediction_Final.docx")
OUTPUT_PDF = Path(
    "/Users/sharanpatil/Downloads/stock price prediction/output/pdf/"
    "Stock_Price_Prediction_reference_style.pdf"
)

A4_WIDTH, A4_HEIGHT = A4
LEFT_MARGIN = 0.9 * inch
RIGHT_MARGIN = 0.9 * inch
TOP_MARGIN = 0.85 * inch
BOTTOM_MARGIN = 0.85 * inch
CONTENT_WIDTH = A4_WIDTH - LEFT_MARGIN - RIGHT_MARGIN


def register_times_new_roman() -> None:
    font_paths = {
        "TimesNewRoman": Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf"),
        "TimesNewRoman-Bold": Path(
            "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
        ),
        "TimesNewRoman-Italic": Path(
            "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"
        ),
        "TimesNewRoman-BoldItalic": Path(
            "/System/Library/Fonts/Supplemental/Times New Roman Bold Italic.ttf"
        ),
    }
    for font_name, font_path in font_paths.items():
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    pdfmetrics.registerFontFamily(
        "TimesNewRoman",
        normal="TimesNewRoman",
        bold="TimesNewRoman-Bold",
        italic="TimesNewRoman-Italic",
        boldItalic="TimesNewRoman-BoldItalic",
    )


def build_styles() -> StyleSheet1:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="BodyReference",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=15,
            leading=22.5,
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyJustify",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=15,
            leading=22.5,
            textColor=colors.black,
            alignment=TA_JUSTIFY,
            spaceBefore=0,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CenteredReference",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=15,
            leading=22.5,
            textColor=colors.black,
            alignment=TA_CENTER,
            spaceBefore=0,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading1Reference",
            parent=styles["Heading1"],
            fontName="TimesNewRoman-Bold",
            fontSize=18,
            leading=24,
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceBefore=10,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading2Reference",
            parent=styles["Heading2"],
            fontName="TimesNewRoman-Bold",
            fontSize=16,
            leading=21,
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CaptionReference",
            parent=styles["Italic"],
            fontName="TimesNewRoman-Italic",
            fontSize=13,
            leading=17,
            textColor=colors.black,
            alignment=TA_CENTER,
            spaceBefore=2,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCHeadingReference",
            parent=styles["Heading1"],
            fontName="TimesNewRoman-Bold",
            fontSize=18,
            leading=24,
            textColor=colors.black,
            alignment=TA_CENTER,
            spaceBefore=0,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCLevel1Reference",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=14,
            leading=20,
            textColor=colors.black,
            leftIndent=10,
            firstLineIndent=0,
            spaceBefore=0,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCLevel2Reference",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=13,
            leading=18,
            textColor=colors.black,
            leftIndent=28,
            firstLineIndent=0,
            spaceBefore=0,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableCellReference",
            parent=styles["BodyText"],
            fontName="TimesNewRoman",
            fontSize=12,
            leading=15,
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=0,
        )
    )
    return styles


def iter_block_items(parent: DocxDocument) -> Iterable[DocxParagraph | DocxTable]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, parent)


def paragraph_has_page_break(paragraph: DocxParagraph) -> bool:
    for run in paragraph.runs:
        for child in run._element.iterchildren():
            if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
                return True
    return False


def extract_inline_images(paragraph: DocxParagraph) -> list[tuple[bytes, float | None, float | None]]:
    images: list[tuple[bytes, float | None, float | None]] = []
    for run in paragraph.runs:
        blips = run._element.xpath('.//*[local-name()="blip"]')
        for blip in blips:
            rel_id = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not rel_id:
                continue
            image_part = paragraph.part.related_parts[rel_id]
            width_pt = None
            height_pt = None
            inline = blip
            while inline is not None and inline.tag not in {qn("wp:inline"), qn("wp:anchor")}:
                inline = inline.getparent()
            if inline is not None:
                extent = inline.find(qn("wp:extent"))
                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    if cx:
                        width_pt = int(cx) / 12700.0
                    if cy:
                        height_pt = int(cy) / 12700.0
            images.append((image_part.blob, width_pt, height_pt))
    return images


def run_markup(run) -> str:
    chunks: list[str] = []
    for child in run._element.iterchildren():
        if child.tag == qn("w:t"):
            chunks.append(escape(child.text or ""))
        elif child.tag == qn("w:tab"):
            chunks.append("    ")
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
                continue
            chunks.append("<br/>")
    text = "".join(chunks)
    if not text:
        return ""
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text


def paragraph_markup(paragraph: DocxParagraph) -> str:
    content = "".join(run_markup(run) for run in paragraph.runs)
    if content:
        return content
    return escape(paragraph.text or "")


def style_for_paragraph(paragraph: DocxParagraph, styles: StyleSheet1) -> ParagraphStyle:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""
    alignment = paragraph.alignment

    if style_name == "Heading 1":
        return styles["Heading1Reference"]
    if style_name == "Heading 2":
        return styles["Heading2Reference"]
    if text.startswith("Fig.") or text.startswith("Figure "):
        return styles["CaptionReference"]
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return styles["CenteredReference"]
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return styles["BodyJustify"]
    return styles["BodyReference"]


def paragraph_to_flowable(
    paragraph: DocxParagraph, styles: StyleSheet1
) -> list[Paragraph | Spacer | Image | PageBreak]:
    flowables: list[Paragraph | Spacer | Image | PageBreak] = []
    text = paragraph.text.strip()
    images = extract_inline_images(paragraph)
    has_page_break = paragraph_has_page_break(paragraph)

    if text:
        markup = paragraph_markup(paragraph)
        if paragraph.style and paragraph.style.name == "List Paragraph":
            markup = f"- {markup}"
        para = Paragraph(markup, style_for_paragraph(paragraph, styles))
        if paragraph.style and paragraph.style.name == "Heading 1":
            para.toc_level = 0
        elif paragraph.style and paragraph.style.name == "Heading 2":
            para.toc_level = 1
        flowables.append(para)
    elif not images:
        flowables.append(Spacer(1, 8))

    for blob, width_pt, height_pt in images:
        image_stream = io.BytesIO(blob)
        image_reader = ImageReader(image_stream)
        pixel_width, pixel_height = image_reader.getSize()
        target_width = width_pt or pixel_width
        target_height = height_pt or pixel_height
        scale = min(CONTENT_WIDTH / target_width, (A4_HEIGHT * 0.55) / target_height, 1.0)
        flowable = Image(
            io.BytesIO(blob),
            width=target_width * scale,
            height=target_height * scale,
        )
        flowable.hAlign = "CENTER"
        flowables.append(flowable)
        flowables.append(Spacer(1, 8))

    if has_page_break:
        flowables.append(PageBreak())

    return flowables


def table_to_flowable(table: DocxTable, styles: StyleSheet1) -> LongTable:
    row_count = len(table.rows)
    col_count = max(len(row.cells) for row in table.rows)
    col_width = CONTENT_WIDTH / max(col_count, 1)
    data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            cell_lines = []
            for paragraph in cell.paragraphs:
                cell_text = paragraph.text.strip()
                if cell_text:
                    cell_lines.append(escape(cell_text))
            if not cell_lines:
                cell_lines = [" "]
            row_cells.append(Paragraph("<br/>".join(cell_lines), styles["TableCellReference"]))
        while len(row_cells) < col_count:
            row_cells.append(Paragraph(" ", styles["TableCellReference"]))
        data.append(row_cells)

    table_flowable = LongTable(
        data,
        colWidths=[col_width] * col_count,
        repeatRows=1 if row_count > 1 else 0,
        hAlign="CENTER",
    )
    table_flowable.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "TimesNewRoman"),
                ("FONTSIZE", (0, 0), (-1, -1), 12),
                ("LEADING", (0, 0), (-1, -1), 15),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table_flowable


class ReferenceDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs) -> None:
        super().__init__(filename, **kwargs)
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="main-frame",
        )
        template = PageTemplate(id="main-template", frames=[frame], onPage=self.draw_footer)
        self.addPageTemplates([template])

    def draw_footer(self, canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("TimesNewRoman", 11)
        canvas.drawCentredString(A4_WIDTH / 2, 0.45 * inch, str(doc.page))
        canvas.restoreState()

    def afterFlowable(self, flowable) -> None:
        if hasattr(flowable, "toc_level"):
            text = flowable.getPlainText()
            self.notify("TOCEntry", (flowable.toc_level, text, self.page))


def build_story(document: DocxDocument, styles: StyleSheet1) -> list:
    story = []
    skip_manual_toc_entries = False
    manual_toc_inserted = False
    for block in iter_block_items(document):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if text.upper() == "TABLE OF CONTENTS":
                skip_manual_toc_entries = True
                manual_toc_inserted = True
                story.append(Paragraph(text, styles["TOCHeadingReference"]))
                toc = TableOfContents()
                toc.levelStyles = [styles["TOCLevel1Reference"], styles["TOCLevel2Reference"]]
                story.append(Spacer(1, 6))
                story.append(toc)
                story.append(PageBreak())
                continue
            if skip_manual_toc_entries:
                if block.style and block.style.name == "Heading 1":
                    skip_manual_toc_entries = False
                else:
                    continue
            story.extend(paragraph_to_flowable(block, styles))
        else:
            if skip_manual_toc_entries:
                continue
            story.append(table_to_flowable(block, styles))
            story.append(Spacer(1, 10))
    if not manual_toc_inserted:
        story.insert(0, Paragraph("Table of Contents", styles["TOCHeadingReference"]))
        toc = TableOfContents()
        toc.levelStyles = [styles["TOCLevel1Reference"], styles["TOCLevel2Reference"]]
        story.insert(1, toc)
        story.insert(2, PageBreak())
    return story


def main() -> None:
    register_times_new_roman()
    styles = build_styles()
    document = Document(str(SOURCE_DOCX))
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    doc = ReferenceDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=LEFT_MARGIN,
        rightMargin=RIGHT_MARGIN,
        topMargin=TOP_MARGIN,
        bottomMargin=BOTTOM_MARGIN,
        title="Stock Price Prediction",
        author="Codex",
    )
    story = build_story(document, styles)
    doc.multiBuild(story)
    print(f"Created PDF: {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
